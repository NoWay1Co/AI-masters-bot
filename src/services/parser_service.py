import httpx
from bs4 import BeautifulSoup
from typing import List, Optional, Dict, Tuple
from urllib.parse import urljoin, urlparse
import re
from ..data.models import Program, Course, ProgramType, ProgramDetails
from ..utils.logger import logger
from datetime import datetime, timedelta
from pypdf import PdfReader
from docx import Document
import openpyxl
from io import BytesIO
from .cache_service import cache_service
from pathlib import Path
from ..utils.config import settings
import json

class ITMOParser:
    def __init__(self):
        self.base_url = "https://abit.itmo.ru"
        self.programs_urls = {
            ProgramType.AI: "https://abit.itmo.ru/program/master/ai",
            ProgramType.AI_PRODUCT: "https://abit.itmo.ru/program/master/ai_product"
        }
        # Создаем папку для файлов
        self.files_dir = settings.DATA_DIR / "files"
        self.files_dir.mkdir(parents=True, exist_ok=True)
    
    async def parse_all_programs(self) -> List[Program]:
        cache_key = "all_programs"
        
        # Проверяем кэш
        cached_programs = await cache_service.get(cache_key)
        if cached_programs:
            logger.info("Using cached programs", count=len(cached_programs))
            return cached_programs
        
        programs = []
        
        for program_type, url in self.programs_urls.items():
            try:
                program = await self._parse_program_page(program_type, url)
                if program:
                    programs.append(program)
            except Exception as e:
                logger.error("Failed to parse program", program_type=program_type, error=str(e))
        
        # Если парсинг не дал результатов, используем mock данные
        if not programs or all(len(p.courses) == 0 for p in programs):
            logger.info("Using mock data for demonstration")
            programs = self._get_mock_programs()
        
        # Сохраняем в кэш на 6 часов
        if programs:
            await cache_service.set(cache_key, programs, timedelta(hours=6))
        
        return programs
    
    def _get_mock_programs(self) -> List[Program]:
        """Временные mock данные для демонстрации функциональности"""
        
        ai_courses = [
            Course(
                id="ai_1",
                name="Основы машинного обучения",
                credits=6,
                hours=216,
                semester=1,
                is_elective=False,
                block="Модули (дисциплины)",
                category="Обязательные дисциплины",
                description="Изучение базовых алгоритмов машинного обучения, линейных моделей, деревьев решений."
            ),
            Course(
                id="ai_2", 
                name="Глубокое обучение",
                credits=6,
                hours=216,
                semester=2,
                is_elective=False,
                block="Модули (дисциплины)",
                category="Обязательные дисциплины",
                description="Нейронные сети, сверточные и рекуррентные архитектуры, трансформеры."
            ),
            Course(
                id="ai_3",
                name="Компьютерное зрение", 
                credits=5,
                hours=180,
                semester=2,
                is_elective=True,
                block="Модули (дисциплины)",
                category="Пул выборных дисциплин",
                description="Обработка изображений, детекция объектов, семантическая сегментация."
            ),
            Course(
                id="ai_4",
                name="Обработка естественного языка",
                credits=5,
                hours=180,
                semester=3,
                is_elective=True,
                block="Модули (дисциплины)",
                category="Пул выборных дисциплин",
                description="Анализ текстов, языковые модели, машинный перевод."
            ),
            Course(
                id="ai_5",
                name="Математические методы в ИИ",
                credits=4,
                hours=144,
                semester=1,
                is_elective=False,
                block="Модули (дисциплины)",
                category="Обязательные дисциплины",
                description="Линейная алгебра, теория вероятностей, оптимизация для ИИ."
            ),
            Course(
                id="ai_6",
                name="Этика ИИ и безопасность",
                credits=3,
                hours=108,
                semester=3,
                is_elective=True,
                block="Модули (дисциплины)",
                category="Пул выборных дисциплин",
                description="Этические аспекты ИИ, безопасность алгоритмов, объяснимость."
            ),
            Course(
                id="ai_7",
                name="Проектная деятельность",
                credits=8,
                hours=288,
                semester=4,
                is_elective=False,
                block="Практики",
                category="Производственная практика",
                description="Выполнение итогового проекта с применением изученных технологий ИИ."
            )
        ]
        
        ai_product_courses = [
            Course(
                id="aip_1",
                name="Управление AI-продуктами",
                credits=6,
                hours=216,
                semester=1,
                is_elective=False,
                block="Модули (дисциплины)",
                category="Обязательные дисциплины",
                description="Методология разработки AI-продуктов, жизненный цикл, метрики."
            ),
            Course(
                id="aip_2",
                name="Data Science для продуктов",
                credits=6,
                hours=216,
                semester=1,
                is_elective=False,
                block="Модули (дисциплины)",
                category="Обязательные дисциплины",
                description="Анализ данных, A/B тестирование, построение аналитических дашбордов."
            ),
            Course(
                id="aip_3",
                name="MLOps и инфраструктура",
                credits=5,
                hours=180,
                semester=2,
                is_elective=False,
                block="Модули (дисциплины)",
                category="Обязательные дисциплины",
                description="Развертывание ML-моделей, мониторинг, CI/CD для ML."
            ),
            Course(
                id="aip_4",
                name="Бизнес-анализ AI-решений",
                credits=4,
                hours=144,
                semester=2,
                is_elective=True,
                block="Модули (дисциплины)",
                category="Пул выборных дисциплин",
                description="ROI AI-проектов, оценка эффективности, бизнес-кейсы."
            ),
            Course(
                id="aip_5",
                name="UX/UI для AI-продуктов",
                credits=4,
                hours=144,
                semester=3,
                is_elective=True,
                block="Модули (дисциплины)",
                category="Пул выборных дисциплин",
                description="Проектирование интерфейсов с ИИ, пользовательский опыт."
            ),
            Course(
                id="aip_6",
                name="Правовые аспекты ИИ",
                credits=3,
                hours=108,
                semester=3,
                is_elective=True,
                block="Модули (дисциплины)",
                category="Пул выборных дисциплин",
                description="Регулирование ИИ, GDPR, интеллектуальная собственность."
            ),
            Course(
                id="aip_7",
                name="Стартап в области ИИ",
                credits=6,
                hours=216,
                semester=4,
                is_elective=False,
                block="Практики",
                category="Производственная практика",
                description="Создание AI-стартапа от идеи до MVP и первых продаж."
            )
        ]
        
        programs = [
            Program(
                id="ai",
                name="Искусственный интеллект",
                type=ProgramType.AI,
                url="https://abit.itmo.ru/program/master/ai",
                courses=ai_courses,
                total_credits=sum(course.credits for course in ai_courses),
                duration_semesters=4,
                description="""Программа готовит специалистов в области искусственного интеллекта. 
                Студенты изучают машинное обучение, нейронные сети, компьютерное зрение, 
                обработку естественного языка. Обучение включает работу над реальными проектами 
                с ведущими IT-компаниями.""",
                details=ProgramDetails(
                    form_of_study="Очная",
                    duration="2 года",
                    language="Русский, Английский",
                    cost_per_year="350 000 рублей",
                    about_program="Магистерская программа по искусственному интеллекту"
                ),
                parsed_at=datetime.now()
            ),
            Program(
                id="ai_product",
                name="Управление ИИ-продуктами/AI Product",
                type=ProgramType.AI_PRODUCT,
                url="https://abit.itmo.ru/program/master/ai_product",
                courses=ai_product_courses,
                total_credits=sum(course.credits for course in ai_product_courses),
                duration_semesters=4,
                description="""Программа для тех, кто хочет создавать и управлять AI-продуктами. 
                Изучение методологий разработки, бизнес-анализа, UX для ИИ, правовых аспектов. 
                Выпускники становятся продакт-менеджерами в области ИИ.""",
                details=ProgramDetails(
                    form_of_study="Очная",
                    duration="2 года", 
                    language="Русский, Английский",
                    cost_per_year="350 000 рублей",
                    about_program="Магистерская программа по управлению ИИ-продуктами"
                ),
                parsed_at=datetime.now()
            )
        ]
        
        return programs
    
    async def _parse_program_page(self, program_type: ProgramType, url: str) -> Optional[Program]:
        async with httpx.AsyncClient(timeout=30.0) as client:
            response = await client.get(url)
            response.raise_for_status()
        
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # Извлекаем название программы
        title_elem = soup.find('h1') or soup.find('title')
        program_name = title_elem.get_text(strip=True) if title_elem else f"Program {program_type}"
        
        # Извлекаем описание и дополнительную информацию
        description = await self._extract_description(soup)
        details = await self._extract_program_details(soup)
        
        # Ищем ссылку на учебный план
        curriculum_url = await self._find_curriculum_link(soup, url)
        
        # Парсим учебный план
        courses = []
        if curriculum_url:
            courses = await self._parse_curriculum_file(curriculum_url, program_type)
        
        # Если курсы не найдены через URL, пробуем локальные файлы
        if not courses:
            local_curricula = await self.load_local_curriculum_files()
            if local_curricula:
                # Берем первый найденный учебный план
                courses = list(local_curricula.values())[0]
        
        program = Program(
            id=program_type.value,
            name=program_name,
            type=program_type,
            url=url,
            courses=courses,
            total_credits=sum(course.credits for course in courses),
            duration_semesters=max((course.semester for course in courses), default=4),
            description=description,
            details=details,
            parsed_at=datetime.now()
        )
        
        logger.info("Program parsed", program_id=program.id, courses_count=len(courses))
        return program
    
    async def _extract_description(self, soup: BeautifulSoup) -> Optional[str]:
        # Ищем описание программы в различных местах
        selectors = [
            '.program-description',
            '.content-main p',
            '.program-info',
            'meta[name="description"]',
            '.hero-description',
            '.program-about'
        ]
        
        for selector in selectors:
            elem = soup.select_one(selector)
            if elem:
                if elem.name == 'meta':
                    return elem.get('content', '').strip()
                else:
                    text = elem.get_text(strip=True)
                    if len(text) > 50:  # Только содержательные описания
                        return text
        
        return None
    
    async def _extract_program_details(self, soup: BeautifulSoup) -> Optional[ProgramDetails]:
        """Извлекает детали программы с очисткой и структурированием данных"""
        details = ProgramDetails()
        
        # Получаем весь текст страницы для анализа
        page_text = soup.get_text()
        
        try:
            # Поиск в JSON данных страницы
            next_data_script = soup.find('script', {'id': '__NEXT_DATA__'})
            if next_data_script and next_data_script.string:
                data = json.loads(next_data_script.string)
                self._extract_details_from_json_clean(data, details)
            
            # Дополнительное извлечение из HTML с регулярными выражениями
            self._extract_details_from_html_clean(soup, page_text, details)
            
        except Exception as e:
            logger.debug("Failed to extract program details", error=str(e))
        
        return details if any(vars(details).values()) else None
    
    def _extract_details_from_json_clean(self, data: dict, details: ProgramDetails):
        """Извлекает детали из JSON данных с очисткой значений"""
        def find_value(obj, keys):
            if isinstance(obj, dict):
                for key, value in obj.items():
                    if key.lower() in keys and isinstance(value, str):
                        return value
                    elif isinstance(value, (dict, list)):
                        result = find_value(value, keys)
                        if result:
                            return result
            elif isinstance(obj, list):
                for item in obj:
                    result = find_value(item, keys)
                    if result:
                        return result
            return None
        
        # Поиск различных полей с последующей очисткой
        raw_form = find_value(data, ['form_of_study', 'форма обучения', 'form'])
        raw_duration = find_value(data, ['duration', 'длительность', 'period'])
        raw_language = find_value(data, ['language', 'язык', 'languages'])
        raw_cost = find_value(data, ['cost', 'стоимость', 'price', 'tuition'])
        raw_dormitory = find_value(data, ['dormitory', 'общежитие', 'housing'])
        raw_military = find_value(data, ['military', 'военный', 'military_center'])
        raw_accreditation = find_value(data, ['accreditation', 'аккредитация'])
        raw_about = find_value(data, ['about', 'о программе', 'description'])
        raw_manager = find_value(data, ['program_manager', 'менеджер программы', 'manager'])
        raw_additional = find_value(data, ['additional_opportunities', 'дополнительные возможности', 'opportunities'])
        raw_directions = find_value(data, ['study_directions', 'направления подготовки', 'directions'])
        
        # Очищаем найденные значения
        if raw_form:
            details.form_of_study = self._clean_form_of_study(raw_form)
        if raw_duration:
            details.duration = self._clean_duration(raw_duration)
        if raw_language:
            details.language = self._clean_language(raw_language)
        if raw_cost:
            details.cost_per_year = self._clean_cost(raw_cost)
        if raw_dormitory:
            details.dormitory = self._clean_yes_no_field(raw_dormitory)
        if raw_military:
            details.military_center = self._clean_yes_no_field(raw_military)
        if raw_accreditation:
            details.accreditation = self._clean_yes_no_field(raw_accreditation)
        if raw_about:
            details.about_program = self._clean_about_program(raw_about)
        if raw_manager:
            manager_name, manager_contacts = self._clean_program_manager(raw_manager)
            details.program_manager = manager_name
            details.manager_contacts = manager_contacts
        if raw_additional:
            details.additional_opportunities = self._clean_additional_opportunities(raw_additional)
        if raw_directions:
            details.study_directions = self._clean_study_directions(raw_directions)
    
    def _extract_details_from_html_clean(self, soup: BeautifulSoup, page_text: str, details: ProgramDetails):
        """Извлекает детали из HTML текста с использованием регулярных выражений"""
        
        # Паттерны для извлечения информации
        patterns = {
            'form_of_study': [
                r'форма\s+обучения[:\s]*([^\n\r]+?)(?:длительность|язык|стоимость|$)',
                r'(очная|заочная|очно-заочная)',
            ],
            'duration': [
                r'длительность[:\s]*([^\n\r]+?)(?:язык|стоимость|общежитие|$)',
                r'(\d+)\s*года?\b',
            ],
            'language': [
                r'язык\s+обучения[:\s]*([^\n\r]+?)(?:стоимость|общежитие|военный|$)',
                r'(русский|английский)',
            ],
            'cost_per_year': [
                r'стоимость\s+контрактного\s+обучения[:\s]*\([^)]*\)[:\s]*([^\n\r]+?)(?:общежитие|военный|гос|$)',
                r'(\d+[\s\d]*)\s*₽',
            ],
            'dormitory': [
                r'общежитие[:\s]*([^\n\r]+?)(?:военный|гос|дополнительные|$)',
                r'общежитие[:\s]*(да|нет)',
            ],
            'military_center': [
                r'военный\s+учебный\s+центр[:\s]*([^\n\r]+?)(?:гос|дополнительные|менеджер|$)',
                r'военный[:\s]*(да|нет)',
            ],
            'accreditation': [
                r'гос\.\s*аккредитация[:\s]*([^\n\r]+?)(?:дополнительные|менеджер|программа|$)',
                r'аккредитация[:\s]*(да|нет)',
            ],
        }
        
        # Применяем паттерны только если поля еще не заполнены
        for field, field_patterns in patterns.items():
            current_value = getattr(details, field)
            if not current_value:  # Заполняем только пустые поля
                for pattern in field_patterns:
                    match = re.search(pattern, page_text, re.I | re.MULTILINE)
                    if match:
                        value = match.group(1) if match.lastindex and match.lastindex >= 1 else match.group(0)
                        cleaned_value = self._clean_field_value(field, value.strip())
                        if cleaned_value:
                            setattr(details, field, cleaned_value)
                            break
        
        # Поиск менеджера программы
        if not details.program_manager:
            manager_patterns = [
                r'менеджер\s+программы[:\s]*([^@\n\r]+)[@]',
                r'менеджер\s+программы[:\s]*([A-ZА-Я][a-zа-я]+\s+[A-ZА-Я][a-zа-я]+\s+[A-ZА-Я][a-zа-я]+)',
            ]
            
            for pattern in manager_patterns:
                match = re.search(pattern, page_text, re.I | re.MULTILINE)
                if match:
                    manager_name, manager_contacts = self._clean_program_manager(match.group(1))
                    details.program_manager = manager_name
                    if manager_contacts:
                        details.manager_contacts = manager_contacts
                    break
        
        # Поиск описания программы
        if not details.about_program:
            # Ищем в HTML структуре по id="about"
            about_header = soup.find('h2', {'id': 'about'})
            if about_header:
                # Собираем весь текст из блока "О программе"
                about_content = []
                
                # Ищем все элементы после заголовка до следующего h2, h3 или определенных секций
                current_elem = about_header.find_next_sibling()
                while current_elem:
                    # Останавливаемся на следующих заголовках или секциях
                    if (current_elem.name in ['h1', 'h2', 'h3'] or 
                        (current_elem.get_text(strip=True).upper() in ['КАРЬЕРА', 'ПАРТНЕРЫ ПРОГРАММЫ', 'КОМАНДА ОБРАЗОВАТЕЛЬНОЙ ПРОГРАММЫ', 'УЧЕБНЫЙ ПЛАН'])):
                        break
                    
                    # Получаем текст из текущего элемента
                    text = current_elem.get_text(strip=True)
                    if text and len(text) > 10:  # Только содержательный текст
                        about_content.append(text)
                    
                    current_elem = current_elem.find_next_sibling()
                
                # Также ищем вложенные элементы в div/section после заголовка
                about_section = about_header.find_next(['div', 'section'])
                if about_section and not about_content:
                    # Ищем все текстовые элементы внутри секции
                    text_elements = about_section.find_all(['p', 'div', 'span'], text=True)
                    for elem in text_elements:
                        text = elem.get_text(strip=True)
                        if text and len(text) > 20:
                            about_content.append(text)
                
                # Объединяем найденный контент
                if about_content:
                    full_text = ' '.join(about_content)
                    details.about_program = self._clean_about_program(full_text)
            
            # Если HTML парсинг не сработал, используем регулярные выражения как fallback
            if not details.about_program:
                about_patterns = [
                    r'о\s+программе[:\s]*([^карьера]{100,}?)(?:карьера|ты\s+сможешь|выпускники|показать\s+все|партнеры|команда|учебный\s+план)',
                    r'создавайте\s+[^\.]*\.\s*([^карьера]{100,}?)(?:карьера|показать\s+все)',
                ]
                
                for pattern in about_patterns:
                    match = re.search(pattern, page_text, re.I | re.MULTILINE | re.DOTALL)
                    if match:
                        details.about_program = self._clean_about_program(match.group(1))
                        break
        

        
        # Поиск дополнительных возможностей
        if not details.additional_opportunities:
            additional_patterns = [
                r'дополнительные\s+возможности[:\s]*([^менеджер]+?)(?:менеджер|программа|$)',
                r'дополнительные\s+возможности[:\s]*([^\.]+\.)',
            ]
            
            for pattern in additional_patterns:
                match = re.search(pattern, page_text, re.I | re.MULTILINE)
                if match:
                    details.additional_opportunities = self._clean_additional_opportunities(match.group(1))
                    break
        
        # Поиск направлений подготовки
        if not details.study_directions:
            directions_patterns = [
                r'направления\s+подготовки[:\s]*([^о\s]+?)(?:о\s+программе|$)',
                r'направления\s+подготовки[:\s]*([^\.]+\.)',
            ]
            
            for pattern in directions_patterns:
                match = re.search(pattern, page_text, re.I | re.MULTILINE)
                if match:
                    details.study_directions = self._clean_study_directions(match.group(1))
                    break
        

        

    
    def _clean_field_value(self, field: str, value: str) -> str:
        """Очищает значение поля в зависимости от его типа"""
        if field == 'form_of_study':
            return self._clean_form_of_study(value)
        elif field == 'duration':
            return self._clean_duration(value)
        elif field == 'language':
            return self._clean_language(value)
        elif field == 'cost_per_year':
            return self._clean_cost(value)
        elif field in ['dormitory', 'military_center', 'accreditation']:
            return self._clean_yes_no_field(value)
        else:
            return value.strip()
    
    def _clean_form_of_study(self, value: str) -> str:
        """Очищает поле 'форма обучения'"""
        value_lower = value.lower()
        if 'очная' in value_lower:
            return 'очная'
        elif 'заочная' in value_lower:
            return 'заочная'
        elif 'очно-заочная' in value_lower:
            return 'очно-заочная'
        return value.strip()
    
    def _clean_duration(self, value: str) -> str:
        """Очищает поле 'длительность'"""
        # Ищем паттерн "2 года"
        match = re.search(r'(\d+)\s*года?', value, re.I)
        if match:
            return f"{match.group(1)} года"
        
        # Если просто число
        match = re.search(r'^(\d+)$', value.strip())
        if match:
            return f"{match.group(1)} года"
            
        return value.strip()
    
    def _clean_language(self, value: str) -> str:
        """Очищает поле 'язык обучения'"""
        value_lower = value.lower()
        languages = []
        if 'русский' in value_lower:
            languages.append('русский')
        if 'английский' in value_lower:
            languages.append('английский')
        
        if languages:
            return ', '.join(languages)
        return value.strip()
    
    def _clean_cost(self, value: str) -> str:
        """Очищает поле 'стоимость'"""
        # Ищем паттерн с числом и рублями
        match = re.search(r'(\d+[\s\d]*)\s*(?:₽|рубл)', value, re.I)
        if match:
            cost = match.group(1).replace(' ', ' ')  # Нормализуем пробелы
            return f"{cost} ₽"
        return value.strip()
    
    def _clean_yes_no_field(self, value: str) -> str:
        """Очищает поля типа да/нет"""
        value_lower = value.lower().strip()
        if value_lower.startswith('да') or 'да' in value_lower[:5]:
            return 'да'
        elif value_lower.startswith('нет') or 'нет' in value_lower[:5]:
            return 'нет'
        return value.strip()
    
    def _clean_program_manager(self, value: str) -> tuple[str, str]:
        """Разделяет менеджера программы на имя и контакты"""
        # Ищем полное имя в русском формате
        name_patterns = [
            r'([А-Я][а-я]+\s+[А-Я][а-я]+\s+[А-Я][а-я]+)',
            r'менеджер\s+программы[:\s]*([А-Я][а-я]+\s+[А-Я][а-я]+\s+[А-Я][а-я]+)',
        ]
        
        name = ""
        for pattern in name_patterns:
            match = re.search(pattern, value)
            if match:
                name = match.group(1).strip()
                break
        
        # Ищем контакты (email и телефон)
        contacts = []
        
        # Поиск email
        email_match = re.search(r'([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})', value)
        if email_match:
            contacts.append(email_match.group(1))
        
        # Поиск телефона
        phone_patterns = [
            r'(\+7\s*\(\d{3}\)\s*\d{3}-\d{2}-\d{2})',
            r'(\+7\s*\d{3}\s*\d{3}-\d{2}-\d{2})',
        ]
        
        for pattern in phone_patterns:
            phone_match = re.search(pattern, value)
            if phone_match:
                contacts.append(phone_match.group(1))
                break
        
        return name, ', '.join(contacts) if contacts else ""
    
    def _clean_about_program(self, value: str) -> str:
        """Очищает описание программы - берет полный текст из блока О ПРОГРАММЕ"""
        if not value:
            return ""
            
        # Удаляем лишние пробелы и переносы
        cleaned = re.sub(r'\s+', ' ', value).strip()
        
        # Убираем технические элементы и навигационные части
        cleaned = re.sub(r'показать\s+все.*$', '', cleaned, flags=re.I)
        cleaned = re.sub(r'партнеры\s+программы.*$', '', cleaned, flags=re.I)
        cleaned = re.sub(r'команда\s+образовательной.*$', '', cleaned, flags=re.I)
        cleaned = re.sub(r'учебный\s+план.*$', '', cleaned, flags=re.I)
        cleaned = re.sub(r'карьера.*$', '', cleaned, flags=re.I)
        cleaned = re.sub(r'вступительные\s+экзамены.*$', '', cleaned, flags=re.I)
        cleaned = re.sub(r'как\s+поступить.*$', '', cleaned, flags=re.I)
        cleaned = re.sub(r'ты\s+сможешь\s+работать.*$', '', cleaned, flags=re.I)
        cleaned = re.sub(r'выпускники\s+программы.*$', '', cleaned, flags=re.I)
        
        # Убираем повторяющиеся фразы и служебные элементы
        cleaned = re.sub(r'подробнее\s*', '', cleaned, flags=re.I)
        cleaned = re.sub(r'смотреть\s+вакансии\s*', '', cleaned, flags=re.I)
        
        # Очищаем и возвращаем текст
        cleaned = cleaned.strip()
        
        # Если текст слишком длинный, обрезаем до разумного размера
        if len(cleaned) > 1000:
            # Обрезаем до последнего полного предложения в пределах 1000 символов
            truncated = cleaned[:1000]
            last_dot = truncated.rfind('.')
            if last_dot > 500:  # Если есть предложение достаточной длины
                cleaned = truncated[:last_dot + 1]
            else:
                cleaned = truncated + '...'
        
        return cleaned
    
    def _clean_additional_opportunities(self, value: str) -> str:
        """Очищает поле дополнительных возможностей"""
        cleaned = re.sub(r'\s+', ' ', value).strip()
        
        # Убираем лишние части
        cleaned = re.sub(r'менеджер\s+программы.*$', '', cleaned, flags=re.I)
        cleaned = re.sub(r'программа в сфере ии.*$', '', cleaned, flags=re.I)
        cleaned = re.sub(r'даты вступительного.*$', '', cleaned, flags=re.I)
        
        # Ищем конкретные возможности в правильном порядке
        opportunities = []
        if 'онлайн' in cleaned.lower():
            opportunities.append('Онлайн')
        if 'трек аспирантуры' in cleaned.lower():
            opportunities.append('Трек аспирантуры')
        if 'пиш' in cleaned.lower():
            opportunities.append('ПИШ')
        if 'соп' in cleaned.lower():
            opportunities.append('СОП')
        
        if opportunities:
            return ', '.join(opportunities)
        
        # Если ничего конкретного не найдено, но есть текст
        if len(cleaned) > 10:
            return cleaned[:50] + '...' if len(cleaned) > 50 else cleaned
        
        return ""
    
    def _clean_study_directions(self, value: str) -> str:
        """Очищает направления подготовки"""
        cleaned = re.sub(r'\s+', ' ', value).strip()
        
        # Ищем полные коды направлений с названиями
        direction_patterns = [
            r'(\d{2}\.\d{2}\.\d{2}\s*[А-Я][^0-9]{20,}?)(?=\d{2}\.\d{2}|\d+бюджетных|$)',
            r'(\d{2}\.\d{2}\.\d{2}[^0-9]+?)(?=\d+бюджетных|$)',
        ]
        
        directions = []
        for pattern in direction_patterns:
            matches = re.findall(pattern, cleaned)
            for match in matches:
                direction = re.sub(r'\s+', ' ', match).strip()
                # Убираем числа в конце (количество мест)
                direction = re.sub(r'\d+бюджетных.*$', '', direction).strip()
                if direction and len(direction) > 15:
                    directions.append(direction)
        
        if directions:
            return '; '.join(directions[:2])  # Максимум 2 направления
        
        return cleaned.strip()

    
    async def _find_curriculum_link(self, soup: BeautifulSoup, base_url: str) -> Optional[str]:
        """
        Ищет ссылку на учебный план, включая поиск в JSON данных страницы
        """
        logger.info("Searching for curriculum link", base_url=base_url)
        
        # 1. Ищем в JSON данных страницы (Next.js __NEXT_DATA__)
        try:
            next_data_script = soup.find('script', {'id': '__NEXT_DATA__'})
            if next_data_script and next_data_script.string:
                data = json.loads(next_data_script.string)
                
                # Ищем academic_plan в JSON
                def find_academic_plan(obj, path=""):
                    if isinstance(obj, dict):
                        for key, value in obj.items():
                            if key == 'academic_plan' and isinstance(value, str) and value.startswith('http'):
                                logger.info("Found academic_plan in JSON", url=value, path=f"{path}.{key}")
                                return value
                            elif isinstance(value, (dict, list)):
                                result = find_academic_plan(value, f"{path}.{key}")
                                if result:
                                    return result
                    elif isinstance(obj, list):
                        for i, item in enumerate(obj):
                            result = find_academic_plan(item, f"{path}[{i}]")
                            if result:
                                return result
                    return None
                
                academic_plan_url = find_academic_plan(data)
                if academic_plan_url:
                    return academic_plan_url
                    
        except Exception as e:
            logger.debug("Failed to parse JSON data", error=str(e))
        
        # 2. Ищем URL в JavaScript коде
        try:
            scripts = soup.find_all('script')
            for script in scripts:
                if script.string:
                    # Ищем URL паттерны для учебных планов
                    patterns = [
                        r'["\']https?://[^"\']*(?:plan|curriculum|academic)[^"\']*\.pdf["\']',
                        r'["\']https?://api\.itmo\.su[^"\']*plan[^"\']*\.pdf["\']',
                        r'academic_plan["\']:\s*["\']([^"\']+\.pdf)["\']'
                    ]
                    
                    for pattern in patterns:
                        matches = re.findall(pattern, script.string, re.I)
                        for match in matches:
                            url = match.strip('\'"')
                            if url.startswith('http'):
                                logger.info("Found curriculum URL in JavaScript", url=url)
                                return url
        except Exception as e:
            logger.debug("Failed to parse JavaScript", error=str(e))
        
        # 3. Традиционный поиск ссылок
        # Ищем кнопку "Скачать учебный план" - теперь и в button элементах
        for tag in ['a', 'button']:
            download_elements = soup.find_all(tag, string=re.compile(r'скачать\s+учебный\s+план', re.I))
            for elem in download_elements:
                href = elem.get('href') or elem.get('data-href') or elem.get('onclick')
                if href and 'http' in href:
                    url_match = re.search(r'https?://[^\s"\']+\.pdf', href)
                    if url_match:
                        logger.info("Found curriculum in button/link", url=url_match.group())
                        return url_match.group()
        
        # 4. Поиск по различным селекторам
        all_links = soup.find_all('a', href=True)
        for link in all_links:
            link_text = link.get_text(strip=True).lower()
            href = link.get('href')
            
            if href and not href.startswith('#'):
                # Проверяем текст ссылки
                if any(word in link_text for word in ['учебный план', 'curriculum', 'план обучения']):
                    full_url = urljoin(base_url, href)
                    logger.info("Found curriculum link by text", url=full_url, text=link_text)
                    return full_url
                
                # Проверяем расширение файла
                if any(ext in href.lower() for ext in ['.pdf', '.docx', '.xlsx']):
                    if any(word in href.lower() for word in ['plan', 'curriculum', 'учебн', 'academic']):
                        full_url = urljoin(base_url, href)
                        logger.info("Found curriculum file by extension", url=full_url)
                        return full_url
        
        logger.warning("No curriculum link found")
        return None
    
    async def _parse_curriculum_file(self, file_url: str, program_type: ProgramType) -> List[Course]:
        cache_key = f"curriculum_{file_url}"
        
        # Проверяем кэш
        cached_courses = await cache_service.get(cache_key)
        if cached_courses:
            logger.info("Using cached curriculum", file_url=file_url, count=len(cached_courses))
            return cached_courses
        
        try:
            async with httpx.AsyncClient(timeout=30.0) as client:
                response = await client.get(file_url)
                response.raise_for_status()
            
            file_content = response.content
            file_extension = self._get_file_extension(file_url)
            
            # Сохраняем файл локально для будущего использования
            await self._save_curriculum_file(file_url, file_content, file_extension, program_type)
            
            courses = []
            if file_extension == '.pdf':
                courses = await self._parse_pdf_curriculum(file_content)
            elif file_extension == '.docx':
                courses = await self._parse_docx_curriculum(file_content)
            elif file_extension == '.xlsx':
                courses = await self._parse_xlsx_curriculum(file_content)
            else:
                logger.warning("Unsupported file format", file_url=file_url)
                return []
            
            # Сохраняем в кэш на 12 часов
            if courses:
                await cache_service.set(cache_key, courses, timedelta(hours=12))
            
            return courses
        
        except Exception as e:
            logger.error("Failed to parse curriculum file", file_url=file_url, error=str(e))
            return []
    
    async def _save_curriculum_file(self, file_url: str, content: bytes, extension: str, program_type: ProgramType) -> str:
        """Сохраняет скачанный файл учебного плана в data/files"""
        try:
            # Создаем имя файла на основе типа программы
            filename_mapping = {
                ProgramType.AI: "10033-abit-3.pdf",
                ProgramType.AI_PRODUCT: "pdf.pdf"
            }
            
            filename = filename_mapping.get(program_type, f"curriculum_{datetime.now().strftime('%Y%m%d_%H%M%S')}")
            
            if not filename.endswith(extension):
                filename += extension
                
            file_path = self.files_dir / filename
            
            # Сохраняем файл
            with open(file_path, 'wb') as f:
                f.write(content)
            
            logger.info("Curriculum file saved", file_path=str(file_path), size=len(content))
            return str(file_path)
            
        except Exception as e:
            logger.error("Failed to save curriculum file", error=str(e))
            return ""
    
    async def load_local_curriculum_files(self) -> Dict[str, List[Course]]:
        """Загружает курсы из локально сохраненных файлов в data/files"""
        local_curricula = {}
        
        if not self.files_dir.exists():
            return local_curricula
            
        for file_path in self.files_dir.glob("*"):
            if file_path.is_file() and file_path.suffix in ['.pdf', '.docx', '.xlsx']:
                try:
                    with open(file_path, 'rb') as f:
                        content = f.read()
                    
                    courses = []
                    if file_path.suffix == '.pdf':
                        courses = await self._parse_pdf_curriculum(content)
                    elif file_path.suffix == '.docx':
                        courses = await self._parse_docx_curriculum(content)
                    elif file_path.suffix == '.xlsx':
                        courses = await self._parse_xlsx_curriculum(content)
                    
                    if courses:
                        local_curricula[file_path.stem] = courses
                        logger.info("Loaded local curriculum", file=file_path.name, courses_count=len(courses))
                        
                except Exception as e:
                    logger.error("Failed to load local curriculum", file=file_path.name, error=str(e))
        
        return local_curricula
    
    def _get_file_extension(self, url: str) -> str:
        """Определяет расширение файла по URL"""
        parsed = urlparse(url)
        path = parsed.path.lower()
        
        # Проверяем стандартные расширения в конце пути
        for ext in ['.pdf', '.docx', '.xlsx']:
            if path.endswith(ext):
                return ext
        
        # Специальная обработка для API ИТМО - ищем тип в пути
        if 'api.itmo.su' in url.lower() and '/plan/' in path:
            if path.endswith('/pdf') or '/pdf/' in path:
                return '.pdf'
            elif path.endswith('/docx') or '/docx/' in path:
                return '.docx'
            elif path.endswith('/xlsx') or '/xlsx/' in path:
                return '.xlsx'
        
        # Проверяем расширения в любом месте URL (fallback)
        url_lower = url.lower()
        if '.pdf' in url_lower:
            return '.pdf'
        elif '.docx' in url_lower:
            return '.docx'
        elif '.xlsx' in url_lower:
            return '.xlsx'
        
        return ''
    
    async def _parse_pdf_curriculum(self, content: bytes) -> List[Course]:
        """Улучшенный парсинг PDF с учетом реальной структуры учебного плана"""
        courses = []
        
        try:
            with BytesIO(content) as pdf_file:
                reader = PdfReader(pdf_file)
                
                full_text = ""
                for page in reader.pages:
                    full_text += page.extract_text() + "\n"
                
                courses = self._extract_courses_from_curriculum_text(full_text)
        
        except Exception as e:
            logger.error("Failed to parse PDF curriculum", error=str(e))
        
        return courses
    
    async def _parse_docx_curriculum(self, content: bytes) -> List[Course]:
        courses = []
        
        try:
            with BytesIO(content) as docx_file:
                doc = Document(docx_file)
                
                full_text = ""
                for paragraph in doc.paragraphs:
                    full_text += paragraph.text + "\n"
                
                # Также парсим таблицы
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(cell.text for cell in row.cells)
                        full_text += row_text + "\n"
                
                courses = self._extract_courses_from_curriculum_text(full_text)
        
        except Exception as e:
            logger.error("Failed to parse DOCX curriculum", error=str(e))
        
        return courses
    
    async def _parse_xlsx_curriculum(self, content: bytes) -> List[Course]:
        courses = []
        
        try:
            with BytesIO(content) as xlsx_file:
                workbook = openpyxl.load_workbook(xlsx_file)
                
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    sheet_courses = self._extract_courses_from_excel_sheet_improved(sheet)
                    courses.extend(sheet_courses)
        
        except Exception as e:
            logger.error("Failed to parse XLSX curriculum", error=str(e))
        
        return courses
    
    def _extract_courses_from_curriculum_text(self, text: str) -> List[Course]:
        """Улучшенное извлечение курсов из текста учебного плана"""
        courses = []
        lines = text.split('\n')
        
        current_semester = 1
        current_category = "Обязательные дисциплины"
        current_block = "Модули (дисциплины)"
        course_id_counter = 0
        
        logger.info("Starting curriculum parsing", total_lines=len(lines))
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
            
            # Определяем текущий блок
            if re.search(r'Блок\s+\d+\.?\s*(.+)', line, re.I):
                match = re.search(r'Блок\s+\d+\.?\s*(.+)', line, re.I)
                current_block = match.group(1).strip()
                logger.debug("Found block", block=current_block)
                continue
            
            # Определяем категорию (обязательные/выборные)
            if re.search(r'Обязательные\s+дисциплины', line, re.I):
                current_category = "Обязательные дисциплины"
                logger.debug("Switched to mandatory courses")
                continue
            elif re.search(r'Пул\s+выборных\s+дисциплин', line, re.I):
                current_category = "Пул выборных дисциплин"
                logger.debug("Switched to elective courses")
                continue
            
            # Определяем семестр
            semester_match = re.search(r'(\d+)\s+семестр', line, re.I)
            if semester_match:
                current_semester = int(semester_match.group(1))
                logger.debug("Found semester", semester=current_semester)
                continue
            
            # Парсим строку с курсом
            course = self._parse_course_line(line, current_semester, current_category, current_block, course_id_counter)
            if course:
                courses.append(course)
                course_id_counter += 1
                logger.debug("Parsed course", name=course.name, credits=course.credits, semester=course.semester)
        
        logger.info("Curriculum parsing completed", total_courses=len(courses))
        return courses
    
    def _parse_course_line(self, line: str, semester: int, category: str, block: str, course_id: int) -> Optional[Course]:
        """Парсит строку с информацией о курсе"""
        
        # Паттерны для различных форматов строк курсов
        patterns = [
            # Семестр + Название + Кредиты + Часы
            r'^(\d+)\s+(.+?)\s+(\d+)\s+(\d+)$',
            # Название + Кредиты + Часы (семестр определен выше)
            r'^(.+?)\s+(\d+)\s+(\d+)$',
            # Семестр + Название (кредиты и часы в другом месте)
            r'^(\d+)\s+(.+?)$'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, line.strip())
            if match:
                groups = match.groups()
                
                # Определяем параметры курса в зависимости от количества групп
                if len(groups) == 4:  # Семестр + Название + Кредиты + Часы
                    try:
                        sem = int(groups[0])
                        name = groups[1].strip()
                        credits = int(groups[2])
                        hours = int(groups[3])
                    except ValueError:
                        continue
                elif len(groups) == 3:  # Название + Кредиты + Часы
                    try:
                        name = groups[0].strip()
                        credits = int(groups[1])
                        hours = int(groups[2])
                        sem = semester
                    except ValueError:
                        continue
                elif len(groups) == 2:  # Семестр + Название или Название + одно число
                    try:
                        if groups[0].isdigit() and not groups[1].isdigit():
                            # Семестр + Название
                            sem = int(groups[0])
                            name = groups[1].strip()
                            credits = 3  # По умолчанию
                            hours = credits * 36  # Стандартный расчет
                        else:
                            # Название + кредиты/часы
                            name = groups[0].strip()
                            credits = int(groups[1]) if groups[1].isdigit() else 3
                            hours = credits * 36
                            sem = semester
                    except ValueError:
                        continue
                else:
                    continue
                
                # Фильтруем нереалистичные значения
                if credits > 50 or hours > 2000:  # Слишком большие значения
                    # Возможно, часы и кредиты перепутаны
                    if hours <= 50:
                        credits, hours = hours, credits
                    else:
                        continue
                
                # Фильтруем слишком короткие или бессмысленные названия
                if len(name) < 3 or name.isdigit() or not any(c.isalpha() for c in name):
                    continue
                
                # Создаем курс
                course = Course(
                    id=f"course_{course_id}",
                    name=name,
                    credits=credits,
                    hours=hours,
                    semester=sem,
                    is_elective="выборн" in category.lower(),
                    block=block,
                    category=category
                )
                
                return course
        
        return None
    
    def _extract_courses_from_excel_sheet_improved(self, sheet) -> List[Course]:
        """Улучшенное извлечение курсов из Excel листа"""
        courses = []
        
        # Ищем заголовки колонок более тщательно
        headers = {}
        header_row = 0
        
        for row_idx, row in enumerate(sheet.iter_rows(max_row=20)):
            for col_idx, cell in enumerate(row):
                if cell.value:
                    value = str(cell.value).lower()
                    
                    if any(word in value for word in ['семестр', 'semester']):
                        headers['semester'] = col_idx
                    elif any(word in value for word in ['наименование', 'дисципли', 'модул', 'название', 'name']):
                        headers['name'] = col_idx
                    elif any(word in value for word in ['з.е', 'кредит', 'зачет', 'credit']):
                        headers['credits'] = col_idx
                    elif any(word in value for word in ['час', 'hour']):
                        headers['hours'] = col_idx
            
            if len(headers) >= 3:  # Нашли основные колонки
                header_row = row_idx
                break
        
        if len(headers) < 2:
            logger.warning("Could not find proper headers in Excel sheet")
            return courses
        
        current_category = "Обязательные дисциплины"
        current_block = "Модули (дисциплины)"
        current_semester = 1
        
        # Парсим данные
        for row_idx, row in enumerate(sheet.iter_rows(min_row=header_row + 2)):
            try:
                # Проверяем на служебные строки
                first_cell = str(row[0].value or "").strip()
                
                # Определяем категорию и блоки
                if re.search(r'Обязательные\s+дисциплины', first_cell, re.I):
                    current_category = "Обязательные дисциплины"
                    continue
                elif re.search(r'Пул\s+выборных\s+дисциплин', first_cell, re.I):
                    current_category = "Пул выборных дисциплин"
                    continue
                elif re.search(r'Блок\s+\d+', first_cell, re.I):
                    match = re.search(r'Блок\s+\d+\.?\s*(.+)', first_cell, re.I)
                    if match:
                        current_block = match.group(1).strip()
                    continue
                
                # Извлекаем данные курса
                name = ""
                credits = 0
                hours = 0
                semester = current_semester
                
                if 'name' in headers and len(row) > headers['name'] and row[headers['name']].value:
                    name = str(row[headers['name']].value).strip()
                
                if 'credits' in headers and len(row) > headers['credits'] and row[headers['credits']].value:
                    credits_val = str(row[headers['credits']].value)
                    credits_match = re.search(r'\d+', credits_val)
                    credits = int(credits_match.group()) if credits_match else 0
                
                if 'hours' in headers and len(row) > headers['hours'] and row[headers['hours']].value:
                    hours_val = str(row[headers['hours']].value)
                    hours_match = re.search(r'\d+', hours_val)
                    hours = int(hours_match.group()) if hours_match else 0
                
                if 'semester' in headers and len(row) > headers['semester'] and row[headers['semester']].value:
                    semester_val = str(row[headers['semester']].value)
                    semester_match = re.search(r'\d+', semester_val)
                    if semester_match:
                        semester = int(semester_match.group())
                        current_semester = semester
                
                # Рассчитываем недостающие значения
                if credits > 0 and hours == 0:
                    hours = credits * 36  # Стандартный коэффициент
                elif hours > 0 and credits == 0:
                    credits = max(1, hours // 36)
                
                # Проверяем валидность данных
                if name and len(name) > 3 and credits > 0 and not name.isdigit():
                    course = Course(
                        id=f"course_{len(courses)}",
                        name=name,
                        credits=credits,
                        hours=hours,
                        semester=semester,
                        is_elective="выборн" in current_category.lower(),
                        block=current_block,
                        category=current_category
                    )
                    courses.append(course)
            
            except Exception as e:
                logger.debug("Failed to parse excel row", row_number=row_idx, error=str(e))
                continue
        
        return courses 